# sistemas/assistencia_judiciaria/core/document.py
import os
import logging
import re
from typing import Union

logger = logging.getLogger("sistemas.assistencia_judiciaria.core.document")

def strip_markdown_markers(text: str) -> str:
    """Remove marcadores markdown do texto"""
    # Remove negrito **texto** -> texto
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    # Remove itálico *texto* -> texto
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    return text

def process_docx_inline_formatting(paragraph, text: str, base_italic: bool = False, clear_paragraph: bool = False):
    """
    Processa formatação inline markdown e adiciona runs formatados ao parágrafo DOCX
    """
    # Limpar o parágrafo apenas se solicitado
    if clear_paragraph:
        paragraph.clear()

    # Processar texto com regex para encontrar formatações
    current_text = ""
    i = 0
    
    while i < len(text):
        # Negrito **texto**
        if text[i:i+2] == '**':
            end_pos = text.find('**', i+2)
            if end_pos != -1:
                if current_text:
                    run = paragraph.add_run(current_text)
                    if base_italic: run.italic = True
                    current_text = ""
                
                bold_text = text[i+2:end_pos]
                run = paragraph.add_run(bold_text)
                run.bold = True
                if base_italic: run.italic = True
                
                i = end_pos + 2
                continue

        # Itálico *texto* (mas não **)
        elif text[i] == '*' and i+1 < len(text) and text[i+1] != '*':
            end_pos = text.find('*', i+1)
            if end_pos != -1 and not (text[i-1:i+1] == '**' if i > 0 else False):
                if current_text:
                    run = paragraph.add_run(current_text)
                    if base_italic: run.italic = True
                    current_text = ""
                
                italic_text = text[i+1:end_pos]
                run = paragraph.add_run(italic_text)
                run.italic = True
                
                i = end_pos + 1
                continue

        # Citações "texto"
        elif text[i] == '"':
            end_pos = text.find('"', i+1)
            if end_pos != -1:
                if current_text:
                    run = paragraph.add_run(current_text)
                    if base_italic: run.italic = True
                    current_text = ""
                
                quote_text = text[i:end_pos+1]
                run = paragraph.add_run(quote_text)
                run.italic = True
                
                i = end_pos + 1
                continue

        # Caractere normal
        current_text += text[i]
        i += 1

    # Adicionar texto restante
    if current_text:
        run = paragraph.add_run(current_text)
        if base_italic:
            run.italic = True

def markdown_to_docx(markdown_text: str, output_path: str, numero_processo: str = "") -> Union[bool, str]:
    """
    Converte markdown para DOCX usando python-docx
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Tentar carregar template personalizado
        template_path = os.path.join("templates", "template.docx")
        if os.path.exists(template_path):
            try:
                doc = Document(template_path)
                # Limpar conteúdo existente
                for p in list(doc.paragraphs):
                    p._element.getparent().remove(p._element)
            except Exception:
                doc = Document()
        else:
            doc = Document()
            # Margens padrão
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)

        lines = markdown_text.split('\n')
        list_counter = 0

        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                list_counter = 0
                continue

            # Cabeçalhos
            if line.startswith('### '):
                p = doc.add_paragraph()
                run = p.add_run(strip_markdown_markers(line[4:]))
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(46, 74, 107)
                list_counter = 0
            elif line.startswith('## '):
                p = doc.add_paragraph()
                run = p.add_run(strip_markdown_markers(line[3:]))
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(46, 74, 107)
                list_counter = 0
            elif line.startswith('# '):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(strip_markdown_markers(line[2:]))
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(46, 74, 107)
                list_counter = 0

            # Listas
            elif line.startswith('- '):
                text = line[2:]
                list_counter += 1
                if list_counter <= 26:
                    marker = f"{chr(ord('a') + list_counter - 1)}) "
                else:
                    marker = "• "
                
                p = doc.add_paragraph()
                p.add_run(marker)
                process_docx_inline_formatting(p, text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Citações
            elif line.startswith('> '):
                text = line[2:]
                p = doc.add_paragraph()
                process_docx_inline_formatting(p, text, base_italic=True)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.left_indent = Inches(0.5)
                p.right_indent = Inches(0.5)
                list_counter = 0

            # Texto normal
            else:
                p = doc.add_paragraph()
                process_docx_inline_formatting(p, line)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                list_counter = 0

        doc.save(output_path)
        return True

    except ImportError:
        return "ERRO: python-docx não instalado."
    except Exception as e:
        logger.exception("Erro ao gerar DOCX")
        return f"ERRO: {str(e)}"

def docx_to_pdf(docx_path: str, pdf_path: str) -> Union[bool, str]:
    """
    Converte DOCX para PDF usando LibreOffice ou docx2pdf
    """
    try:
        import subprocess
        
        # Opção 1: LibreOffice
        try:
            output_dir = os.path.dirname(pdf_path)
            cmd = [
                "soffice", "--headless", "--convert-to", "pdf",
                "--outdir", output_dir, docx_path
            ]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            if result.returncode == 0:
                docx_name = os.path.splitext(os.path.basename(docx_path))[0]
                generated_pdf = os.path.join(output_dir, docx_name + ".pdf")
                if os.path.exists(generated_pdf) and generated_pdf != pdf_path:
                    if os.path.exists(pdf_path): os.remove(pdf_path)
                    os.rename(generated_pdf, pdf_path)
                return True
        except Exception:
            pass

        # Opção 2: docx2pdf
        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            return True
        except ImportError:
            pass

        return "ERRO: Instale LibreOffice ou docx2pdf para gerar PDF."

    except Exception as e:
        return f"ERRO: {str(e)}"
