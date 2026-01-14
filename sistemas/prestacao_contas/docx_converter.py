# sistemas/prestacao_contas/docx_converter.py
"""
Conversor de Parecer para DOCX

Converte o parecer de prestação de contas em documento Word formatado.

Autor: LAB/PGE-MS
"""

import os
import re
from datetime import datetime
from pathlib import Path
from typing import Optional, List

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Diretório do módulo
MODULE_DIR = Path(__file__).parent

# Diretório para arquivos temporários
TEMP_DIR = MODULE_DIR / "temp_docs"
TEMP_DIR.mkdir(exist_ok=True)

# Caminho da imagem do logo para cabeçalho
LOGO_PATH = MODULE_DIR.parent.parent / "logo" / "logo-pge.png"

# Meses em português
MESES_PT_BR = {
    1: "janeiro",
    2: "fevereiro",
    3: "março",
    4: "abril",
    5: "maio",
    6: "junho",
    7: "julho",
    8: "agosto",
    9: "setembro",
    10: "outubro",
    11: "novembro",
    12: "dezembro"
}


def _criar_documento_base() -> Document:
    """Cria documento com configurações base"""
    doc = Document()

    # Configura margens ABNT
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # Adiciona cabeçalho com logo e rodapé institucional
    _add_header_footer(doc)

    # Adiciona borda de página
    _add_page_border(doc)

    return doc


def _add_header_footer(doc: Document):
    """
    Adiciona cabeçalho e rodapé padrão com imagem do logo.

    Cabeçalho: Logo da PGE centralizado
    Rodapé: Texto institucional centralizado
    """
    for section in doc.sections:
        # Configura distância do cabeçalho e rodapé
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.0)

        # ===== CABEÇALHO =====
        header = section.header
        header.is_linked_to_previous = False

        # Limpa parágrafos existentes e adiciona um novo
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.clear()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Remove recuos para centralização perfeita
        header_para.paragraph_format.left_indent = Cm(0)
        header_para.paragraph_format.right_indent = Cm(0)
        header_para.paragraph_format.first_line_indent = Cm(0)

        # Adiciona imagem do logo no cabeçalho
        if LOGO_PATH.exists():
            run = header_para.add_run()
            run.add_picture(str(LOGO_PATH), width=Cm(5.6))  # Largura de 5.6cm
        else:
            # Fallback: texto se não houver logo
            run = header_para.add_run("PROCURADORIA-GERAL DO ESTADO DE MATO GROSSO DO SUL")
            run.font.name = "Arial"
            run.font.size = Pt(10)
            run.font.bold = True

        # ===== RODAPÉ =====
        footer = section.footer
        footer.is_linked_to_previous = False

        # Limpa parágrafos existentes e adiciona um novo
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.clear()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Texto do rodapé
        run = footer_para.add_run("Procuradoria-Geral do Estado de Mato Grosso do Sul")
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.italic = True

        # Adiciona linha separadora acima do rodapé
        footer_para.paragraph_format.space_before = Pt(6)


def _add_page_border(doc: Document):
    """
    Adiciona borda de página ao documento.
    """
    for section in doc.sections:
        sectPr = section._sectPr

        # Cria elemento de borda de página
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')

        # Define as bordas (superior, inferior, esquerda, direita)
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Espessura da linha
            border.set(qn('w:space'), '24')  # Espaço da borda até o conteúdo
            border.set(qn('w:color'), '000000')  # Cor preta
            pgBorders.append(border)

        sectPr.append(pgBorders)


def _adicionar_cabecalho(doc: Document, numero_cnj: str):
    """Adiciona cabeçalho do parecer"""
    # Título
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("PARECER DE PRESTAÇÃO DE CONTAS")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Arial"

    # Número do processo
    processo = doc.add_paragraph()
    processo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = processo.add_run(f"Processo nº {numero_cnj}")
    run.font.size = Pt(12)
    run.font.name = "Arial"

    # Linha em branco
    doc.add_paragraph()


def _adicionar_parecer_badge(doc: Document, parecer: str):
    """Adiciona indicador visual do parecer"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if parecer == "favoravel":
        texto = "PARECER: FAVORÁVEL"
    elif parecer == "desfavoravel":
        texto = "PARECER: DESFAVORÁVEL"
    else:
        texto = "PARECER: PENDENTE DE ESCLARECIMENTOS"

    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Arial"

    # Linha em branco
    doc.add_paragraph()


def _processar_markdown(doc: Document, texto: str):
    """Processa texto markdown e adiciona ao documento"""
    linhas = texto.split('\n')

    for linha in linhas:
        linha = linha.strip()

        if not linha:
            continue

        # Cabeçalhos
        if linha.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(linha[4:])
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Arial"
            continue

        if linha.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(linha[3:])
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = "Arial"
            continue

        if linha.startswith('# '):
            p = doc.add_paragraph()
            run = p.add_run(linha[2:])
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Arial"
            continue

        # Lista com bullet
        if linha.startswith('- ') or linha.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            texto_item = linha[2:]
            _adicionar_texto_formatado(p, texto_item)
            continue

        # Lista numerada
        match = re.match(r'^(\d+)\.\s+(.+)', linha)
        if match:
            p = doc.add_paragraph(style='List Number')
            texto_item = match.group(2)
            _adicionar_texto_formatado(p, texto_item)
            continue

        # Parágrafo normal
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _adicionar_texto_formatado(p, linha)


def _adicionar_texto_formatado(paragrafo, texto: str):
    """Adiciona texto com formatação inline (negrito, itálico)"""
    # Processa negrito
    partes = re.split(r'\*\*(.+?)\*\*', texto)

    for i, parte in enumerate(partes):
        if not parte:
            continue

        # Partes ímpares são negrito
        if i % 2 == 1:
            run = paragrafo.add_run(parte)
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(12)
        else:
            # Processa itálico dentro do texto normal
            subpartes = re.split(r'\*(.+?)\*', parte)
            for j, subparte in enumerate(subpartes):
                if not subparte:
                    continue
                run = paragrafo.add_run(subparte)
                run.font.name = "Arial"
                run.font.size = Pt(12)
                if j % 2 == 1:
                    run.italic = True


def _adicionar_irregularidades(doc: Document, irregularidades: List[str]):
    """Adiciona seção de irregularidades"""
    p = doc.add_paragraph()
    run = p.add_run("IRREGULARIDADES IDENTIFICADAS")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Arial"

    for irreg in irregularidades:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(irreg)
        run.font.name = "Arial"
        run.font.size = Pt(12)

    doc.add_paragraph()


def _adicionar_rodape(doc: Document):
    """Adiciona rodapé com data em português"""
    doc.add_paragraph()

    # Data em português
    agora = datetime.now()
    mes_pt = MESES_PT_BR[agora.month]
    data = f"{agora.day} de {mes_pt} de {agora.year}"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Campo Grande/MS, {data}")
    run.font.name = "Arial"
    run.font.size = Pt(12)


async def converter_parecer_docx(
    numero_cnj: str,
    parecer: str,
    fundamentacao: str,
    irregularidades: Optional[List[str]] = None,
) -> str:
    """
    Converte parecer para DOCX.

    Args:
        numero_cnj: Número do processo
        parecer: Tipo do parecer (favoravel, desfavoravel, duvida)
        fundamentacao: Texto em markdown da fundamentação
        irregularidades: Lista de irregularidades (se desfavorável)

    Returns:
        Caminho do arquivo DOCX gerado
    """
    doc = _criar_documento_base()

    # Cabeçalho
    _adicionar_cabecalho(doc, numero_cnj)

    # Badge do parecer
    _adicionar_parecer_badge(doc, parecer)

    # Fundamentação (conteúdo direto, sem título)
    _processar_markdown(doc, fundamentacao)

    # Irregularidades (se houver)
    if irregularidades:
        doc.add_paragraph()
        _adicionar_irregularidades(doc, irregularidades)

    # Rodapé
    _adicionar_rodape(doc)

    # Salva documento
    numero_limpo = re.sub(r'[^\d]', '', numero_cnj)
    nome_arquivo = f"parecer_{numero_limpo}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    caminho = TEMP_DIR / nome_arquivo

    doc.save(str(caminho))

    return str(caminho)
