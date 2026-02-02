# sistemas/matriculas_confrontantes/docx_converter.py
"""
Conversor de relatórios de matrículas confrontantes para DOCX.

Reutiliza o DocxConverter do sistema de geração de peças para manter
consistência visual (cabeçalho, rodapé, formatação) em todos os documentos
gerados pelo Portal PGE-MS.

Autor: Portal PGE-MS
"""

import os
import re
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

# Importa o conversor base do gerador de peças
from sistemas.gerador_pecas.docx_converter import DocxConverter, TEMPLATE_PATH, LOGO_PATH


# Diretório do módulo
MODULE_DIR = Path(__file__).parent


class MatriculasDocxConverter(DocxConverter):
    """
    Conversor especializado para relatórios de matrículas confrontantes.

    Herda do DocxConverter do gerador de peças para manter:
    - Mesmo cabeçalho com logo da PGE-MS
    - Mesmo rodapé institucional
    - Mesma formatação base (Times New Roman 12pt, margens ABNT)

    Personalizações para relatórios de matrículas:
    - Metadados sem recuo no início
    - Título centralizado com espaçamento
    """

    def __init__(
        self,
        template_path: Optional[str] = None,
        # Configurações herdadas do DocxConverter
        font_name: str = "Times New Roman",
        font_size: int = 12,
        first_line_indent_cm: float = 2.0,
        line_spacing: float = 1.5,
        space_after_pt: int = 6,
        quote_indent_cm: float = 3.0,
        quote_font_size: int = 11,
        quote_line_spacing: float = 1.0,
        list_indent_cm: float = 3.0,
        margin_top_cm: float = 3.0,
        margin_bottom_cm: float = 2.0,
        margin_left_cm: float = 3.0,
        margin_right_cm: float = 2.0,
        # Linhas em branco após direcionamento (não usado em relatórios)
        linhas_apos_direcionamento: int = 0,
        # Numeração de títulos
        numerar_titulos: bool = False,
    ):
        """
        Inicializa o conversor de relatórios de matrículas.

        Usa o template do gerador de peças por padrão para manter
        consistência visual nos documentos do portal.
        """
        # Usa o template do gerador de peças se não especificado
        if template_path is None and TEMPLATE_PATH.exists():
            template_path = str(TEMPLATE_PATH)

        super().__init__(
            template_path=template_path,
            font_name=font_name,
            font_size=font_size,
            first_line_indent_cm=first_line_indent_cm,
            line_spacing=line_spacing,
            space_after_pt=space_after_pt,
            quote_indent_cm=quote_indent_cm,
            quote_font_size=quote_font_size,
            quote_line_spacing=quote_line_spacing,
            list_indent_cm=list_indent_cm,
            margin_top_cm=margin_top_cm,
            margin_bottom_cm=margin_bottom_cm,
            margin_left_cm=margin_left_cm,
            margin_right_cm=margin_right_cm,
            linhas_apos_direcionamento=linhas_apos_direcionamento,
            numerar_titulos=numerar_titulos,
        )

    def _add_metadata_line(self, doc: Document, label: str, value: str):
        """
        Adiciona linha de metadado sem recuo.
        Formato: **Label:** Valor
        """
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.first_line_indent = Cm(0)  # SEM recuo
        p.paragraph_format.left_indent = Cm(0)

        # Label em negrito
        run_label = p.add_run(f"{label}: ")
        run_label.font.name = self.font_name
        run_label.font.size = Pt(self.font_size)
        run_label.bold = True

        # Valor normal
        run_value = p.add_run(value)
        run_value.font.name = self.font_name
        run_value.font.size = Pt(self.font_size)

    def _add_centered_title(self, doc: Document, text: str):
        """
        Adiciona título centralizado com linha em branco antes e depois.
        """
        # Linha em branco antes
        blank_before = doc.add_paragraph()
        blank_before.paragraph_format.space_before = Pt(0)
        blank_before.paragraph_format.space_after = Pt(0)
        blank_before.paragraph_format.line_spacing = 1.0

        # Título centralizado
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0)

        run = p.add_run(text.upper())
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size)
        run.bold = True

        # Linha em branco depois
        blank_after = doc.add_paragraph()
        blank_after.paragraph_format.space_before = Pt(0)
        blank_after.paragraph_format.space_after = Pt(0)
        blank_after.paragraph_format.line_spacing = 1.0

    def convert_relatorio(
        self,
        relatorio_texto: str,
        output_path: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> bool:
        """
        Converte relatório de matrículas para DOCX com formatação específica.

        Args:
            relatorio_texto: Texto do relatório em Markdown
            output_path: Caminho de saída para o arquivo .docx
            metadata: Metadados opcionais (matricula_principal, data, etc.)

        Returns:
            True se conversão bem sucedida
        """
        try:
            # Reseta contadores de numeração
            self._reset_heading_counters()

            # Carrega template ou cria documento novo
            if self.template_path and os.path.exists(self.template_path):
                doc = Document(self.template_path)
                # Limpa conteúdo existente do template, preservando sectPr (seções)
                for element in doc.element.body[:]:
                    if element.tag.endswith('sectPr'):
                        continue
                    doc.element.body.remove(element)
                self._add_header_footer(doc)
            else:
                doc = Document()
                self._configure_document(doc)

            # 1. Adiciona metadados sem recuo
            if metadata:
                if metadata.get("matricula_principal"):
                    self._add_metadata_line(doc, "Matrícula Principal", str(metadata['matricula_principal']))

                if metadata.get("arquivo"):
                    self._add_metadata_line(doc, "Arquivo", metadata['arquivo'])

                if metadata.get("data_analise"):
                    self._add_metadata_line(doc, "Data da Análise", metadata['data_analise'])

                if metadata.get("modelo"):
                    self._add_metadata_line(doc, "Modelo IA", metadata['modelo'])

            # 2. Adiciona título centralizado com espaçamento
            self._add_centered_title(doc, "RELATÓRIO COMPLETO DO IMÓVEL")

            # 3. Remove o título do relatório se já existir no texto (evita duplicação)
            relatorio_texto = re.sub(
                r'^#\s*RELATÓRIO\s+COMPLETO\s+DO\s+IMÓVEL\s*\n*',
                '',
                relatorio_texto,
                flags=re.IGNORECASE | re.MULTILINE
            )

            # 4. Processa o resto do markdown normalmente
            self._process_markdown(doc, relatorio_texto)

            # Salva o documento
            doc.save(output_path)
            return True

        except Exception as e:
            print(f"Erro na conversão MD→DOCX: {e}")
            import traceback
            traceback.print_exc()
            return False


def gerar_relatorio_docx(
    relatorio_texto: str,
    output_path: str,
    matricula_principal: Optional[str] = None,
    arquivo: Optional[str] = None,
    modelo: Optional[str] = None,
    **kwargs
) -> bool:
    """
    Função de conveniência para gerar relatório DOCX de matrículas.

    Args:
        relatorio_texto: Texto do relatório em Markdown
        output_path: Caminho de saída para o arquivo .docx
        matricula_principal: Número da matrícula principal (opcional)
        arquivo: Nome do arquivo analisado (opcional)
        modelo: Modelo de IA utilizado (opcional)
        **kwargs: Configurações adicionais para o conversor

    Returns:
        True se conversão bem sucedida

    Example:
        >>> gerar_relatorio_docx(
        ...     "# Relatório\\n\\nConteúdo do relatório...",
        ...     "relatorio.docx",
        ...     matricula_principal="12345",
        ...     arquivo="matricula.pdf"
        ... )
    """
    converter = MatriculasDocxConverter(**kwargs)

    # Monta metadados
    metadata = {}
    if matricula_principal:
        metadata["matricula_principal"] = matricula_principal
    if arquivo:
        metadata["arquivo"] = arquivo
    if modelo:
        metadata["modelo"] = modelo

    # Adiciona data/hora atual
    metadata["data_analise"] = datetime.now().strftime("%d/%m/%Y %H:%M")

    return converter.convert_relatorio(
        relatorio_texto=relatorio_texto,
        output_path=output_path,
        metadata=metadata if metadata else None
    )
