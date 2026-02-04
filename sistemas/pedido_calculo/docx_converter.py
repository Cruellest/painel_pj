# sistemas/pedido_calculo/docx_converter.py
"""
Conversor robusto de Markdown para DOCX para Pedido de Cálculo.

Este módulo converte texto em formato Markdown para documentos Word (.docx)
preservando toda a formatação: negrito, itálico, títulos, listas, citações.

REGRAS DE FORMATAÇÃO (Pedido de Cálculo PGE-MS):
- SEM recuo em NENHUMA linha (diferença do gerador de peças)
- Títulos (## headers): SEM recuo, justificados (não centralizados), negrito
- Parágrafos normais: SEM recuo
- Citações (blockquote): Recuo de 4cm, fonte 11pt, itálico
- NÃO gera linha horizontal para '---' (ignorado)

Autor: Portal PGE-MS
"""

import os
import re
from typing import Optional, List, Tuple
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Diretório do módulo
MODULE_DIR = Path(__file__).parent

# Caminho do template DOCX personalizado
TEMPLATE_PATH = MODULE_DIR / "template_pedido_calculo.docx"

# Caminho da imagem do logo para cabeçalho
LOGO_PATH = MODULE_DIR.parent.parent / "logo" / "logo-pge.png"


class DocxConverter:
    """
    Conversor de Markdown para DOCX para Pedido de Cálculo.
    
    Configurações padrão (podem ser sobrescritas pelo template):
    - Fonte: Arial 12pt
    - SEM recuo em NENHUMA linha
    - Espaçamento entre linhas: 1.5
    - Margens: ABNT (3cm esq/sup, 2cm dir/inf)
    - Citações: Recuo 4cm, fonte 11pt, itálico
    - Headers: justificados, SEM recuo, negrito
    """
    
    def __init__(
        self,
        template_path: Optional[str] = None,
        # Configurações de fonte
        font_name: str = "Arial",
        font_size: int = 12,
        # Configurações de parágrafo
        line_spacing: float = 1.5,
        space_after_pt: int = 6,
        # Configurações de citação (blockquote)
        quote_indent_cm: float = 4.0,
        quote_font_size: int = 11,
        quote_line_spacing: float = 1.0,
        # Margens ABNT
        margin_top_cm: float = 3.0,
        margin_bottom_cm: float = 2.0,
        margin_left_cm: float = 3.0,
        margin_right_cm: float = 2.0,
        # Linhas em branco após direcionamento
        linhas_apos_direcionamento: int = 5,
    ):
        """
        Inicializa o conversor com configurações personalizadas.
        """
        self.template_path = template_path or (
            str(TEMPLATE_PATH) if TEMPLATE_PATH.exists() else None
        )
        
        # Configurações de fonte
        self.font_name = font_name
        self.font_size = font_size
        
        # Configurações de parágrafo - SEM RECUO
        self.first_line_indent = Cm(0)  # SEM RECUO
        self.line_spacing = line_spacing
        self.space_after = Pt(space_after_pt)
        
        # Configurações de citação
        self.quote_indent = Cm(quote_indent_cm)
        self.quote_font_size = quote_font_size
        self.quote_line_spacing = quote_line_spacing
        
        # Linhas após direcionamento
        self.linhas_apos_direcionamento = linhas_apos_direcionamento
        
        # Margens
        self.margin_top = Cm(margin_top_cm)
        self.margin_bottom = Cm(margin_bottom_cm)
        self.margin_left = Cm(margin_left_cm)
        self.margin_right = Cm(margin_right_cm)
        
        self._last_was_heading = False
    
    def _clean_markdown(self, markdown_text: str) -> str:
        """
        Limpa e normaliza o markdown antes de processar.
        Corrige problemas comuns de formatação gerados pela IA.
        """
        text = markdown_text
        
        # Corrige padrões problemáticos como "**: **TEXTO" → ": TEXTO" 
        text = re.sub(r'\*\*:\s*\*\*\s*', ': ', text)
        
        # Corrige "**LABEL: **" → "LABEL: " (asteriscos mal fechados)
        text = re.sub(r'\*\*([^*:]+):\s*\*\*\s*', r'\1: ', text)
        
        # Corrige "** texto" → "**texto" (espaço após abertura)
        text = re.sub(r'\*\*\s+(\S)', r'**\1', text)
        
        # Corrige "texto **" → "texto**" (espaço antes de fechamento)
        text = re.sub(r'(\S)\s+\*\*', r'\1**', text)
        
        # Remove asteriscos duplos órfãos que não têm par
        # Padrão: "** " no início de linha ou " **" no final
        text = re.sub(r'^\*\*\s+', '', text, flags=re.MULTILINE)
        text = re.sub(r'\s+\*\*$', '', text, flags=re.MULTILINE)
        
        # Remove "**" sozinho no meio do texto (sem conteúdo entre eles)
        text = re.sub(r'\*\*\*\*', '', text)
        
        return text
    
    def convert(self, markdown_text: str, output_path: str) -> bool:
        """
        Converte texto Markdown para DOCX.
        
        Args:
            markdown_text: Texto em formato Markdown
            output_path: Caminho de saída para o arquivo .docx
            
        Returns:
            True se conversão bem sucedida
        """
        try:
            # Limpa o markdown antes de processar
            markdown_text = self._clean_markdown(markdown_text)
            
            # Carrega template ou cria documento novo
            if self.template_path and os.path.exists(self.template_path):
                doc = Document(self.template_path)
                # Limpa conteúdo existente do template, preservando sectPr (seções)
                for element in doc.element.body[:]:
                    if element.tag.endswith('sectPr'):
                        continue
                    doc.element.body.remove(element)
                self._add_header_footer(doc)
            else:
                doc = Document()
                self._configure_document(doc)
            
            # Processa o Markdown
            self._process_markdown(doc, markdown_text)
            
            # Salva o documento
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"❌ Erro na conversão MD→DOCX: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _configure_document(self, doc: Document):
        """Configura documento com estilos padrão."""
        # Configura margens
        for section in doc.sections:
            section.top_margin = self.margin_top
            section.bottom_margin = self.margin_bottom
            section.left_margin = self.margin_left
            section.right_margin = self.margin_right
        
        # Configura estilo Normal - SEM RECUO
        style = doc.styles['Normal']
        style.font.name = self.font_name
        style.font.size = Pt(self.font_size)
        style.paragraph_format.line_spacing = self.line_spacing
        style.paragraph_format.space_after = self.space_after
        style.paragraph_format.first_line_indent = Cm(0)  # SEM RECUO
        
        # Configura fonte para caracteres asiáticos (evita fallback)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
        
        # Adiciona cabeçalho e rodapé padrão
        self._add_header_footer(doc)
    
    def _add_header_footer(self, doc: Document):
        """
        Adiciona cabeçalho e rodapé padrão com imagem do logo.
        """
        for section in doc.sections:
            section.header_distance = Cm(1.5)
            section.footer_distance = Cm(1.0)
            
            # ===== CABEÇALHO =====
            header = section.header
            header.is_linked_to_previous = False
            
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.clear()
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_para.paragraph_format.left_indent = Cm(0)
            header_para.paragraph_format.right_indent = Cm(0)
            header_para.paragraph_format.first_line_indent = Cm(0)
            
            if LOGO_PATH.exists():
                run = header_para.add_run()
                run.add_picture(str(LOGO_PATH), width=Cm(5.6))
            else:
                run = header_para.add_run("PROCURADORIA-GERAL DO ESTADO DE MATO GROSSO DO SUL")
                run.font.name = self.font_name
                run.font.size = Pt(10)
                run.font.bold = True
            
            # ===== RODAPÉ =====
            footer = section.footer
            footer.is_linked_to_previous = False
            
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.clear()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = footer_para.add_run("Procuradoria-Geral do Estado de Mato Grosso do Sul")
            run.font.name = self.font_name
            run.font.size = Pt(9)
            run.font.italic = True
            
            footer_para.paragraph_format.space_before = Pt(6)
    
    def _process_markdown(self, doc: Document, markdown_text: str):
        """Processa texto Markdown e adiciona ao documento."""
        lines = markdown_text.split('\n')
        i = 0
        in_blockquote = False
        blockquote_lines = []
        is_first_paragraph = True
        found_direcionamento = False
        
        while i < len(lines):
            line_raw = lines[i]
            line = line_raw.strip()
            
            # Blockquote (citação) - pode ter múltiplas linhas
            if line.startswith('> '):
                in_blockquote = True
                blockquote_lines.append(line[2:])
                i += 1
                continue
            elif line.startswith('>'):
                in_blockquote = True
                blockquote_lines.append(line.lstrip('> '))
                i += 1
                continue
            elif in_blockquote and line:
                self._add_blockquote(doc, blockquote_lines)
                blockquote_lines = []
                in_blockquote = False
            elif in_blockquote:
                self._add_blockquote(doc, blockquote_lines)
                blockquote_lines = []
                in_blockquote = False
            
            # Linha vazia
            if not line:
                i += 1
                continue
            
            # Títulos (## headers) - justificados, sem recuo, negrito
            if line.startswith('# '):
                self._add_heading(doc, line[2:], level=1)
                is_first_paragraph = False
            elif line.startswith('## '):
                self._add_heading(doc, line[3:], level=2)
                is_first_paragraph = False
            elif line.startswith('### '):
                self._add_heading(doc, line[4:], level=3)
                is_first_paragraph = False
            elif line.startswith('#### '):
                self._add_heading(doc, line[5:], level=4)
                is_first_paragraph = False
            
            # Linha horizontal ou underscores - IGNORADA
            elif line == '---' or line == '***' or re.match(r'^[_\-]{3,}$', line):
                pass
            
            # Lista não ordenada
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                list_items = []
                while i < len(lines) and (
                    lines[i].strip().startswith('- ') or 
                    lines[i].strip().startswith('* ')
                ):
                    item_text = lines[i].strip()[2:]
                    list_items.append(item_text)
                    i += 1
                self._add_list(doc, list_items, ordered=False)
                is_first_paragraph = False
                continue
            
            # Lista ordenada
            elif re.match(r'^\d+\.\s', line.strip()):
                list_items = []
                while i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
                    item_text = re.sub(r'^\d+\.\s', '', lines[i].strip())
                    list_items.append(item_text)
                    i += 1
                self._add_list(doc, list_items, ordered=True)
                is_first_paragraph = False
                continue
            
            # Lista com letras (a), b), c))
            elif re.match(r'^[a-z]\)\s', line.strip()):
                list_items = []
                while i < len(lines) and re.match(r'^[a-z]\)\s', lines[i].strip()):
                    item_text = re.sub(r'^[a-z]\)\s', '', lines[i].strip())
                    list_items.append(item_text)
                    i += 1
                self._add_list(doc, list_items, ordered=True, style='letter')
                is_first_paragraph = False
                continue
            
            # Parágrafo normal - SEM RECUO
            else:
                line_lower = line.lower()
                line_upper = line.upper()
                
                # Detecta se é o direcionamento
                if is_first_paragraph and ('EXCELENTÍSSIMO' in line_upper or 'EXMO' in line_upper):
                    self._add_direcionamento(doc, line)
                    found_direcionamento = True
                    is_first_paragraph = False
                    
                # Detecta número do processo
                elif self._is_process_number_field(line):
                    if found_direcionamento:
                        for _ in range(self.linhas_apos_direcionamento):
                            blank_para = doc.add_paragraph()
                            blank_para.paragraph_format.space_before = Pt(0)
                            blank_para.paragraph_format.space_after = Pt(0)
                            blank_para.paragraph_format.line_spacing = 1.0
                            blank_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                        found_direcionamento = False
                    self._add_header_field(doc, line, is_last=False)
                    
                # Detecta campos de cabeçalho
                elif self._is_header_field(line):
                    is_last = self._is_last_header_field(line)
                    self._add_header_field(doc, line, is_last=is_last)
                        
                # Parágrafos normais - SEM recuo
                else:
                    self._add_paragraph(doc, line)
                    is_first_paragraph = False
            
            i += 1
        
        # Processa blockquote pendente
        if blockquote_lines:
            self._add_blockquote(doc, blockquote_lines)
    
    def _is_process_number_field(self, line: str) -> bool:
        """Verifica se a linha é um campo de número do processo."""
        line_clean = self._strip_markdown(line).strip().lower()
        
        if not line_clean:
            return False
        
        has_number_indicator = 'nº' in line_clean or 'n°' in line_clean or 'n.' in line_clean
        
        if not has_number_indicator:
            return False
        
        process_patterns = [
            r'^processo',
            r'^autos',
            r'^agravo',
            r'^apelação',
            r'^recurso',
            r'^ação',
            r'^mandado\s+de\s+segurança',
            r'^habeas\s+corpus',
            r'^embargos',
            r'^execução',
            r'^cumprimento\s+de\s+sentença',
            r'^reclamação',
            r'^conflito',
            r'^incidente',
            r'^procedimento',
            r'^petição',
            r'^representação',
            r'^inquérito',
            r'^adi',
            r'^adc',
            r'^adpf',
            r'^processo\s+de\s+origem',
            r'^origem',
        ]
        
        for pattern in process_patterns:
            if re.match(pattern, line_clean):
                return True
        
        return False
    
    def _is_header_field(self, line: str) -> bool:
        """Verifica se a linha é um campo de cabeçalho."""
        line_clean = self._strip_markdown(line).strip().lower()
        
        if not line_clean:
            return False
        
        header_patterns = [
            r'^requerente',
            r'^requerido',
            r'^autor[a]?[\s:\(]',
            r'^réu',
            r'^ré[\s:\(]',
            r'^impetrad[oa]',
            r'^impetrante',
            r'^recorrente',
            r'^recorrido',
            r'^apelante',
            r'^apelado',
            r'^agravante',
            r'^agravado',
            r'^embargante',
            r'^embargado',
            r'^autos\s*n[º°]',
        ]
        
        for pattern in header_patterns:
            if re.match(pattern, line_clean):
                return True
        
        return False
    
    def _is_last_header_field(self, line: str) -> bool:
        """Verifica se é o último campo do cabeçalho."""
        line_clean = self._strip_markdown(line).strip().lower()
        last_patterns = [
            r'^requerido',
            r'^réu',
            r'^ré[\s:\(]',
            r'^embargado',
            r'^apelado',
            r'^agravado',
            r'^recorrido',
            r'^impetrado',
            r'^executado',
            r'^reclamado',
            r'^demandado',
        ]
        for pattern in last_patterns:
            if re.match(pattern, line_clean):
                return True
        return False
    
    def _add_direcionamento(self, doc: Document, text: str):
        """Adiciona direcionamento (primeira linha) sem recuo, espaçamento simples."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0)
        
        clean_text = self._strip_markdown(text).upper()
        
        run = p.add_run(clean_text)
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size)
        run.bold = True
    
    def _add_header_field(self, doc: Document, text: str, is_last: bool = False):
        """Adiciona campo de cabeçalho (Processo nº, Requerente, Requerido)."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0)
        
        # Remove markdown do texto
        clean_text = self._strip_markdown(text)
        
        # Separa etiqueta do valor quando há dois pontos
        if ':' in clean_text:
            parts = clean_text.split(':', 1)
            label = parts[0].strip()
            value = parts[1].strip() if len(parts) > 1 else ''
            
            # Etiqueta em negrito e maiúsculo
            run_label = p.add_run(label.upper() + ':')
            run_label.font.name = self.font_name
            run_label.font.size = Pt(self.font_size)
            run_label.bold = True
            
            # Valor normal
            if value:
                run_value = p.add_run(' ' + value)
                run_value.font.name = self.font_name
                run_value.font.size = Pt(self.font_size)
        else:
            # Se não tem dois pontos, usa formatação normal
            run = p.add_run(clean_text)
            run.font.name = self.font_name
            run.font.size = Pt(self.font_size)
        
        # Se for o último campo, pula uma linha após
        if is_last:
            blank_para = doc.add_paragraph()
            blank_para.paragraph_format.space_before = Pt(0)
            blank_para.paragraph_format.space_after = Pt(0)
            blank_para.paragraph_format.line_spacing = 1.0
            blank_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    
    def _add_heading(self, doc: Document, text: str, level: int):
        """Adiciona título formatado - justificado, sem recuo, negrito."""
        p = doc.add_paragraph()
        
        clean_text = self._strip_markdown(text)
        
        titulo_texto = clean_text.upper() if level <= 2 else clean_text
        
        run = p.add_run(titulo_texto)
        run.font.name = self.font_name
        run.font.bold = True
        run.font.size = Pt(12)
        
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.line_spacing = self.line_spacing
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        
        self._last_was_heading = True
    
    def _add_paragraph(self, doc: Document, text: str):
        """Adiciona parágrafo SEM recuo."""
        p = doc.add_paragraph()
        
        p.style = doc.styles['Normal']
        
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = self.line_spacing
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.space_after = self.space_after
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.right_indent = Cm(0)
        p.paragraph_format.first_line_indent = Cm(0)  # SEM RECUO
        
        self._add_formatted_text(p, text)
        
        self._last_was_heading = False
    
    def _add_formatted_text(self, paragraph, text: str):
        """Adiciona texto com formatação inline (negrito, itálico)."""
        patterns = [
            (r'\*\*\*(.+?)\*\*\*', 'bold_italic'),
            (r'\*\*(.+?)\*\*', 'bold'),
            (r'\*(.+?)\*', 'italic'),
            (r'_(.+?)_', 'italic'),
        ]
        
        segments = []
        last_end = 0
        
        combined_pattern = r'(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*[^*]+?\*|_.+?_)'
        
        for match in re.finditer(combined_pattern, text):
            if match.start() > last_end:
                segments.append(('normal', text[last_end:match.start()]))
            
            matched_text = match.group(0)
            
            if matched_text.startswith('***'):
                clean_text = matched_text[3:-3]
                segments.append(('bold_italic', clean_text))
            elif matched_text.startswith('**'):
                clean_text = matched_text[2:-2]
                segments.append(('bold', clean_text))
            elif matched_text.startswith('*') or matched_text.startswith('_'):
                clean_text = matched_text[1:-1]
                segments.append(('italic', clean_text))
            
            last_end = match.end()
        
        if last_end < len(text):
            segments.append(('normal', text[last_end:]))
        
        if not segments:
            segments = [('normal', text)]
        
        for style, segment_text in segments:
            run = paragraph.add_run(segment_text)
            run.font.name = self.font_name
            run.font.size = Pt(self.font_size)
            
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
            
            if style == 'bold' or style == 'bold_italic':
                run.bold = True
            if style == 'italic' or style == 'bold_italic':
                run.italic = True
    
    def _add_blockquote(self, doc: Document, lines: List[str]):
        """Adiciona citação (blockquote) com formatação especial."""
        text = ' '.join(lines)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = self.quote_indent
        p.paragraph_format.right_indent = Cm(1)
        p.paragraph_format.line_spacing = self.quote_line_spacing
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Cm(0)
        
        self._add_blockquote_text(p, text)
    
    def _add_blockquote_text(self, paragraph, text: str):
        """Adiciona texto de citação com formatação."""
        text = re.sub(r'^\*(.+)\*$', r'\1', text.strip())
        
        pattern = r'\*\*(.+?)\*\*'
        
        last_end = 0
        for match in re.finditer(pattern, text):
            if match.start() > last_end:
                run = paragraph.add_run(text[last_end:match.start()])
                run.font.name = self.font_name
                run.font.size = Pt(self.quote_font_size)
                run.italic = True
            
            run = paragraph.add_run(match.group(1))
            run.font.name = self.font_name
            run.font.size = Pt(self.quote_font_size)
            run.bold = True
            run.italic = True
            
            last_end = match.end()
        
        if last_end < len(text):
            run = paragraph.add_run(text[last_end:])
            run.font.name = self.font_name
            run.font.size = Pt(self.quote_font_size)
            run.italic = True
    
    def _add_list(self, doc: Document, items: List[str], ordered: bool = False, style: str = 'number'):
        """Adiciona lista ordenada ou não ordenada."""
        for i, item in enumerate(items):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Cm(1.25)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = self.line_spacing
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            
            if ordered:
                if style == 'letter':
                    marker = f"{chr(ord('a') + i)}) "
                else:
                    marker = f"{i + 1}. "
            else:
                marker = "• "
            
            run = p.add_run(marker)
            run.font.name = self.font_name
            run.font.size = Pt(self.font_size)
            
            self._add_formatted_text(p, item)
    
    def _strip_markdown(self, text: str) -> str:
        """Remove marcadores markdown do texto."""
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'_(.+?)_', r'\1', text)
        return text


def pedido_calculo_to_docx(
    markdown_text: str,
    output_path: str,
    template_path: Optional[str] = None,
    **kwargs
) -> bool:
    """
    Função de conveniência para converter Pedido de Cálculo de Markdown para DOCX.
    
    Args:
        markdown_text: Texto em formato Markdown
        output_path: Caminho de saída para o arquivo .docx
        template_path: Caminho para template personalizado (opcional)
        **kwargs: Configurações adicionais para DocxConverter
        
    Returns:
        True se conversão bem sucedida
    """
    converter = DocxConverter(template_path=template_path, **kwargs)
    return converter.convert(markdown_text, output_path)
