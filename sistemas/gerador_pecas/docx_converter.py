# sistemas/gerador_pecas/docx_converter.py
"""
Conversor robusto de Markdown para DOCX com template personalizado.

Este módulo converte texto em formato Markdown para documentos Word (.docx)
preservando toda a formatação: negrito, itálico, títulos, listas, citações.

REGRAS DE FORMATAÇÃO (Peças Jurídicas PGE-MS):
- Direcionamento (primeira linha em negrito): SEM recuo, justificado
- Número do processo: SEM recuo, justificado
- 5 linhas em branco entre direcionamento e número do processo
- Títulos (## headers): SEM recuo, justificados (não centralizados), negrito
- Parágrafos normais: COM recuo de 2cm na primeira linha
- Citações (blockquote): Recuo de 3cm, fonte 11pt, SEM itálico
- Listas (bullets/numeração): Recuo de 3cm
- NÃO gera linha horizontal para '---' (ignorado)

Autor: Portal PGE-MS
"""

import os
import re
from typing import Optional, List, Tuple
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Diretório do módulo
MODULE_DIR = Path(__file__).parent

# Caminho do template DOCX personalizado
TEMPLATE_PATH = MODULE_DIR / "template_peca.docx"

# Caminho da imagem do logo para cabeçalho
LOGO_PATH = MODULE_DIR.parent.parent / "logo" / "logo-pge.png"


class DocxConverter:
    """
    Conversor de Markdown para DOCX com suporte a template personalizado.

    Configurações padrão (podem ser sobrescritas pelo template):
    - Fonte: Times New Roman 12pt
    - Recuo primeira linha: 2 cm (apenas parágrafos normais)
    - Espaçamento entre linhas: 1.5
    - Margens: ABNT (3cm esq/sup, 2cm dir/inf)
    - Citações: Recuo 3cm, fonte 11pt, SEM itálico
    - Listas: Recuo 3cm
    - Headers: justificados, SEM recuo, negrito
    """
    
    def __init__(
        self,
        template_path: Optional[str] = None,
        # Configurações de fonte
        font_name: str = "Times New Roman",
        font_size: int = 12,
        # Configurações de parágrafo
        first_line_indent_cm: float = 2.0,
        line_spacing: float = 1.5,
        space_after_pt: int = 6,
        # Configurações de citação (blockquote)
        quote_indent_cm: float = 3.0,
        quote_font_size: int = 11,
        quote_line_spacing: float = 1.0,
        # Configurações de listas
        list_indent_cm: float = 3.0,
        # Margens ABNT
        margin_top_cm: float = 3.0,
        margin_bottom_cm: float = 2.0,
        margin_left_cm: float = 3.0,
        margin_right_cm: float = 2.0,
        # Linhas em branco após direcionamento
        linhas_apos_direcionamento: int = 5,
        # Numeração de títulos
        numerar_titulos: bool = False,  # Desabilitado - IA já gera títulos numerados
    ):
        """
        Inicializa o conversor com configurações personalizadas.
        
        Args:
            template_path: Caminho para template .docx personalizado (opcional)
            font_name: Nome da fonte padrão
            font_size: Tamanho da fonte em pontos
            first_line_indent_cm: Recuo da primeira linha em cm
            line_spacing: Espaçamento entre linhas (1.0, 1.5, 2.0)
            space_after_pt: Espaço após parágrafo em pontos
            quote_indent_cm: Recuo para citações em cm
            quote_font_size: Tamanho da fonte para citações
            quote_line_spacing: Espaçamento entre linhas para citações
            margin_*: Margens do documento em cm
            linhas_apos_direcionamento: Linhas em branco após o direcionamento
            numerar_titulos: Se True, numera títulos automaticamente (1., 2., 2.1., etc.)
        """
        self.template_path = template_path or (
            str(TEMPLATE_PATH) if TEMPLATE_PATH.exists() else None
        )
        
        # Configurações de fonte
        self.font_name = font_name
        self.font_size = font_size
        
        # Configurações de parágrafo
        self.first_line_indent = Cm(first_line_indent_cm)
        self.line_spacing = line_spacing
        self.space_after = Pt(space_after_pt)
        
        # Configurações de citação
        self.quote_indent = Cm(quote_indent_cm)
        self.quote_font_size = quote_font_size
        self.quote_line_spacing = quote_line_spacing

        # Configurações de listas
        self.list_indent = Cm(list_indent_cm)
        
        # Linhas após direcionamento
        self.linhas_apos_direcionamento = linhas_apos_direcionamento
        
        # Margens
        self.margin_top = Cm(margin_top_cm)
        self.margin_bottom = Cm(margin_bottom_cm)
        self.margin_left = Cm(margin_left_cm)
        self.margin_right = Cm(margin_right_cm)
        
        # Numeração de títulos
        self.numerar_titulos = numerar_titulos
        self._heading_counters = [0, 0, 0, 0]  # Contadores para níveis 1, 2, 3, 4
        self._last_was_heading = False  # Flag para controlar parágrafos após headings
    
    def _reset_heading_counters(self):
        """Reseta os contadores de numeração de títulos."""
        self._heading_counters = [0, 0, 0, 0]
        self._last_was_heading = False
    
    def _get_heading_number(self, level: int) -> str:
        """
        Retorna o número do título com base no nível.
        
        Exemplos:
        - Nível 1: 1., 2., 3.
        - Nível 2: 1.1., 1.2., 2.1.
        - Nível 3: 1.1.1., 1.1.2.
        """
        if not self.numerar_titulos:
            return ""
        
        # Ajusta índice (level 1 = index 0)
        idx = level - 1
        if idx < 0 or idx >= len(self._heading_counters):
            return ""
        
        # Incrementa o contador do nível atual
        self._heading_counters[idx] += 1
        
        # Reseta contadores de níveis inferiores
        for i in range(idx + 1, len(self._heading_counters)):
            self._heading_counters[i] = 0
        
        # Monta a numeração
        parts = []
        for i in range(idx + 1):
            if self._heading_counters[i] > 0:
                parts.append(str(self._heading_counters[i]))
        
        return ".".join(parts) + ". " if parts else ""
    
    def convert(self, markdown_text: str, output_path: str) -> bool:
        """
        Converte texto Markdown para DOCX.
        
        Args:
            markdown_text: Texto em formato Markdown
            output_path: Caminho de saída para o arquivo .docx
            
        Returns:
            True se conversão bem sucedida
        """
        try:
            # Reseta contadores de numeração
            self._reset_heading_counters()
            
            # Carrega template ou cria documento novo
            if self.template_path and os.path.exists(self.template_path):
                doc = Document(self.template_path)
                # Limpa conteúdo existente do template, preservando sectPr (seções)
                for element in doc.element.body[:]:
                    # Preserva sectPr para manter cabeçalho/rodapé
                    if element.tag.endswith('sectPr'):
                        continue
                    doc.element.body.remove(element)
                # Adiciona cabeçalho e rodapé (mesmo usando template)
                self._add_header_footer(doc)
            else:
                doc = Document()
                self._configure_document(doc)
            
            # Processa o Markdown
            self._process_markdown(doc, markdown_text)
            
            # Salva o documento
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"❌ Erro na conversão MD→DOCX: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _configure_document(self, doc: Document):
        """Configura documento com estilos padrão."""
        # Configura margens
        for section in doc.sections:
            section.top_margin = self.margin_top
            section.bottom_margin = self.margin_bottom
            section.left_margin = self.margin_left
            section.right_margin = self.margin_right
        
        # Configura estilo Normal
        style = doc.styles['Normal']
        style.font.name = self.font_name
        style.font.size = Pt(self.font_size)
        style.paragraph_format.line_spacing = self.line_spacing
        style.paragraph_format.space_after = self.space_after
        
        # Configura fonte para caracteres asiáticos (evita fallback)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
        
        # Adiciona cabeçalho e rodapé padrão
        self._add_header_footer(doc)
    
    def _add_header_footer(self, doc: Document):
        """
        Adiciona cabeçalho e rodapé padrão com imagem do logo.
        
        Cabeçalho: Logo da PGE centralizado
        Rodapé: Texto institucional centralizado
        """
        for section in doc.sections:
            # Configura distância do cabeçalho e rodapé
            section.header_distance = Cm(1.5)
            section.footer_distance = Cm(1.0)
            
            # ===== CABEÇALHO =====
            header = section.header
            header.is_linked_to_previous = False
            
            # Limpa parágrafos existentes e adiciona um novo
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            header_para.clear()
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Remove recuos para centralização perfeita
            header_para.paragraph_format.left_indent = Cm(0)
            header_para.paragraph_format.right_indent = Cm(0)
            header_para.paragraph_format.first_line_indent = Cm(0)
            
            # Adiciona imagem do logo no cabeçalho (30% menor)
            if LOGO_PATH.exists():
                run = header_para.add_run()
                run.add_picture(str(LOGO_PATH), width=Cm(5.6))  # Largura de 5.6cm (30% menor que 8cm)
            else:
                # Fallback: texto se não houver logo
                run = header_para.add_run("PROCURADORIA-GERAL DO ESTADO DE MATO GROSSO DO SUL")
                run.font.name = self.font_name
                run.font.size = Pt(10)
                run.font.bold = True
            
            # ===== RODAPÉ =====
            footer = section.footer
            footer.is_linked_to_previous = False
            
            # Limpa parágrafos existentes e adiciona um novo
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.clear()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Texto do rodapé
            run = footer_para.add_run("Procuradoria-Geral do Estado de Mato Grosso do Sul")
            run.font.name = self.font_name
            run.font.size = Pt(9)
            run.font.italic = True
            
            # Adiciona linha separadora acima do rodapé (opcional)
            footer_para.paragraph_format.space_before = Pt(6)
    
    def _process_markdown(self, doc: Document, markdown_text: str):
        """Processa texto Markdown e adiciona ao documento."""
        # Remove tags <br> e substitui por quebras de linha reais
        markdown_text = re.sub(r'<br\s*/?>', '\n', markdown_text, flags=re.IGNORECASE)

        lines = markdown_text.split('\n')
        i = 0
        in_blockquote = False
        blockquote_lines = []
        is_first_paragraph = True  # Para detectar direcionamento
        found_direcionamento = False
        
        while i < len(lines):
            line_raw = lines[i]
            line = line_raw.strip()  # Linha sem espaços para detecção
            
            # Blockquote (citação) - pode ter múltiplas linhas
            if line.startswith('> '):
                in_blockquote = True
                blockquote_lines.append(line[2:])
                i += 1
                continue
            elif line.startswith('>'):
                in_blockquote = True
                blockquote_lines.append(line.lstrip('> '))
                i += 1
                continue
            elif in_blockquote and line:
                # Fim do blockquote (linha não começa com >)
                self._add_blockquote(doc, blockquote_lines)
                blockquote_lines = []
                in_blockquote = False
                # Não incrementa i, processa a linha atual
                # (continua para processar esta linha normalmente)
            elif in_blockquote:
                # Linha vazia - fim do blockquote
                self._add_blockquote(doc, blockquote_lines)
                blockquote_lines = []
                in_blockquote = False
            
            # Linha vazia
            if not line:
                i += 1
                continue
            
            # Títulos (## headers) - justificados, sem recuo, negrito
            if line.startswith('# '):
                self._add_heading(doc, line[2:], level=1)
                is_first_paragraph = False
            elif line.startswith('## '):
                self._add_heading(doc, line[3:], level=2)
                is_first_paragraph = False
            elif line.startswith('### '):
                self._add_heading(doc, line[4:], level=3)
                is_first_paragraph = False
            elif line.startswith('#### '):
                self._add_heading(doc, line[5:], level=4)
                is_first_paragraph = False
            
            # Linha horizontal ou underscores - IGNORADA (não gera nada)
            elif line == '---' or line == '***' or re.match(r'^[_\-]{3,}$', line):
                # Ignora completamente - não adiciona nada ao documento
                pass
            
            # Lista não ordenada
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                # Coleta todos os itens da lista
                list_items = []
                while i < len(lines) and (
                    lines[i].strip().startswith('- ') or 
                    lines[i].strip().startswith('* ')
                ):
                    item_text = lines[i].strip()[2:]
                    list_items.append(item_text)
                    i += 1
                self._add_list(doc, list_items, ordered=False)
                is_first_paragraph = False
                continue
            
            # Lista ordenada
            elif re.match(r'^\d+\.\s', line.strip()):
                list_items = []
                while i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
                    item_text = re.sub(r'^\d+\.\s', '', lines[i].strip())
                    list_items.append(item_text)
                    i += 1
                self._add_list(doc, list_items, ordered=True)
                is_first_paragraph = False
                continue
            
            # Lista com letras (a), b), c))
            elif re.match(r'^[a-z]\)\s', line.strip()):
                list_items = []
                while i < len(lines) and re.match(r'^[a-z]\)\s', lines[i].strip()):
                    item_text = re.sub(r'^[a-z]\)\s', '', lines[i].strip())
                    list_items.append(item_text)
                    i += 1
                self._add_list(doc, list_items, ordered=True, style='letter')
                is_first_paragraph = False
                continue
            
            # Parágrafo normal
            else:
                line_lower = line.lower()
                line_upper = line.upper()
                
                # Detecta se é o direcionamento (primeira linha com EXCELENTÍSSIMO, EXMO, À CÂMARA, AO JUIZ, etc.)
                # Pode ou não ter marcação ** de negrito
                if is_first_paragraph and self._is_direcionamento(line_upper):
                    self._add_direcionamento(doc, line)
                    found_direcionamento = True
                    is_first_paragraph = False
                    
                # Detecta número do processo ou tipo de ação com número (Processo nº, Agravo nº, Apelação nº, etc.)
                elif self._is_process_number_field(line):
                    # Adiciona linhas em branco antes do número do processo (se logo após direcionamento)
                    if found_direcionamento:
                        for _ in range(self.linhas_apos_direcionamento):
                            blank_para = doc.add_paragraph()
                            blank_para.paragraph_format.space_before = Pt(0)
                            blank_para.paragraph_format.space_after = Pt(0)
                            blank_para.paragraph_format.line_spacing = 1.0
                            blank_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                        found_direcionamento = False
                    self._add_header_field(doc, line, is_last=False)  # Campo de cabeçalho
                    
                # Detecta campos de cabeçalho (Requerente, Requerido, Autor, Réu) - espaçamento simples
                elif self._is_header_field(line):
                    # Verifica se é o último campo (Requerido/Réu) para pular linha após
                    is_last = self._is_last_header_field(line)
                    self._add_header_field(doc, line, is_last=is_last)
                        
                # Parágrafos normais - COM recuo
                else:
                    self._add_paragraph(doc, line, indent=True)  # Com recuo normal
                    is_first_paragraph = False
            
            i += 1
        
        # Processa blockquote pendente
        if blockquote_lines:
            self._add_blockquote(doc, blockquote_lines)
    
    def _is_process_number_field(self, line: str) -> bool:
        """
        Verifica se a linha é um campo de número do processo ou tipo de ação com número.
        Exemplos: "Processo nº 123", "Agravo de Instrumento nº: 123", "Apelação nº 456", etc.
        """
        # Remove marcação markdown (negrito, itálico) para detectar corretamente
        line_clean = self._strip_markdown(line).strip().lower()
        
        if not line_clean:
            return False
        
        # Verifica se contém indicador de número (nº, n°, n.)
        has_number_indicator = 'nº' in line_clean or 'n°' in line_clean or 'n.' in line_clean
        
        if not has_number_indicator:
            return False
        
        # Padrões de tipos de processo/ação que indicam campo de cabeçalho
        process_patterns = [
            r'^processo',
            r'^autos',
            r'^agravo',
            r'^apelação',
            r'^recurso',
            r'^ação',
            r'^mandado\s+de\s+segurança',
            r'^habeas\s+corpus',
            r'^embargos',
            r'^execução',
            r'^cumprimento\s+de\s+sentença',
            r'^reclamação',
            r'^conflito',
            r'^incidente',
            r'^procedimento',
            r'^petição',
            r'^representação',
            r'^inquérito',
            r'^adi',  # Ação Direta de Inconstitucionalidade
            r'^adc',  # Ação Declaratória de Constitucionalidade
            r'^adpf', # Arguição de Descumprimento de Preceito Fundamental
            r'^processo\s+de\s+origem',  # Processo de Origem nº
            r'^origem',  # Origem nº
        ]
        
        for pattern in process_patterns:
            if re.match(pattern, line_clean):
                return True
        
        return False
    
    def _is_header_field(self, line: str) -> bool:
        """Verifica se a linha é um campo de cabeçalho (Requerente, Requerido, etc.)
        NOTA: Processo nº tem tratamento especial e não está incluído aqui."""
        # Remove marcação markdown (negrito, itálico) para detectar corretamente
        line_clean = self._strip_markdown(line).strip().lower()
        
        # Se a linha estiver vazia, não é campo de cabeçalho
        if not line_clean:
            return False
        
        # Padrões de campos de cabeçalho (SEM processo - tem tratamento separado)
        # Usa regex para capturar variações como "Requerido(s):", "Réu(s):", etc.
        header_patterns = [
            r'^requerente',
            r'^requerido',
            r'^autor[a]?[\s:\(]',
            r'^réu',
            r'^ré[\s:\(]',
            r'^impetrad[oa]',
            r'^impetrante',
            r'^recorrente',
            r'^recorrido',
            r'^apelante',
            r'^apelado',
            r'^agravante',
            r'^interessado',
            r'^agravado',
            r'^embargante',
            r'^embargado',
            r'^autos\s*n[º°]',
        ]
        
        for pattern in header_patterns:
            if re.match(pattern, line_clean):
                return True
        
        return False
    
    def _is_last_header_field(self, line: str) -> bool:
        """Verifica se é o último campo do cabeçalho (parte passiva) para pular linha após."""
        # Remove marcação markdown para detectar corretamente
        line_clean = self._strip_markdown(line).strip().lower()
        # Padrões que indicam fim do cabeçalho (parte passiva em diferentes tipos de ação)
        last_patterns = [
            r'^requerido',
            r'^réu',
            r'^ré[\s:\(]',
            r'^embargado',
            r'^apelado',
            r'^agravado',
            r'^recorrido',
            r'^impetrado',
            r'^executado',
            r'^reclamado',
            r'^demandado',
            r'^interessado',
        ]
        for pattern in last_patterns:
            if re.match(pattern, line_clean):
                return True
        return False

    def _is_direcionamento(self, line_upper: str) -> bool:
        """
        Verifica se a linha é um direcionamento (endereçamento da peça).
        Exemplos: "EXCELENTÍSSIMO SENHOR JUIZ...", "À CÂMARA CÍVEL...", "AO JUÍZO...", etc.
        """
        # Padrões de direcionamento
        direcionamento_patterns = [
            'EXCELENTÍSSIMO',
            'EXCELENTISSIMO',
            'EXMO',
            'À CÂMARA',
            'A CÂMARA',
            'À CAMARA',
            'A CAMARA',
            'AO JUIZ',
            'AO JUÍZO',
            'AO JUIZO',
            'À TURMA',
            'A TURMA',
            'AO TRIBUNAL',
            'À SEÇÃO',
            'A SEÇÃO',
            'À SECAO',
            'A SECAO',
            'AO DESEMBARGADOR',
            'À VARA',
            'A VARA',
            'AO MINISTRO',
            'AO SUPREMO',
            'AO SUPERIOR',
            'À CORTE',
            'A CORTE',
        ]

        for pattern in direcionamento_patterns:
            if pattern in line_upper:
                return True

        return False

    def _add_direcionamento(self, doc: Document, text: str):
        """Adiciona direcionamento (primeira linha) sem recuo, espaçamento simples.
        O direcionamento deve ficar em NEGRITO e MAIÚSCULO."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.0  # Espaçamento simples
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.space_after = Pt(0)  # Sem espaço após (linhas serão adicionadas)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.first_line_indent = Cm(0)  # SEM recuo
        p.paragraph_format.left_indent = Cm(0)  # SEM recuo esquerdo
        
        # Remove markdown do texto e deixa em maiúsculo
        clean_text = self._strip_markdown(text).upper()
        
        # Adiciona como negrito
        run = p.add_run(clean_text)
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size)
        run.bold = True
    
    def _add_header_field(self, doc: Document, text: str, is_last: bool = False):
        """
        Adiciona campo de cabeçalho (Processo nº, Requerente, Requerido).
        Formatação: alinhado à esquerda, espaçamento simples (1.0), sem recuo,
        sem espaçamento entre parágrafos (space_before e space_after = 0).
        A etiqueta (antes dos dois pontos ou do valor) fica em negrito e maiúsculo.
        """
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Alinhado à esquerda (não justificado)
        p.paragraph_format.line_spacing = 1.0  # Espaçamento simples
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.space_after = Pt(0)  # Sem espaço após (campos ficam juntos)
        p.paragraph_format.space_before = Pt(0)  # Sem espaço antes
        p.paragraph_format.first_line_indent = Cm(0)  # SEM recuo
        p.paragraph_format.left_indent = Cm(0)  # SEM recuo esquerdo
        
        # Limpa marcadores markdown órfãos do texto
        text = self._clean_orphan_markers(text)

        # Padrão especial para "Processo nº XXX" (sem dois pontos)
        text_lower = text.lower()
        if 'processo' in text_lower and ('nº' in text_lower or 'n°' in text_lower):
            # Encontra onde termina "nº" ou "n°"
            match = re.search(r'(processo\s*n[º°]\s*:?)', text, re.IGNORECASE)
            if match:
                label = self._strip_markdown(match.group(1).rstrip(':').strip())
                value = self._strip_markdown(text[match.end():].strip().lstrip(':').strip())

                # Etiqueta em negrito e maiúsculo
                run_label = p.add_run(label.upper() + ':')
                run_label.font.name = self.font_name
                run_label.font.size = Pt(self.font_size)
                run_label.bold = True

                # Valor
                if value:
                    run_value = p.add_run(' ' + value)
                    run_value.font.name = self.font_name
                    run_value.font.size = Pt(self.font_size)
            else:
                # Fallback
                self._add_formatted_text(p, text)
        # Separa etiqueta do valor quando há dois pontos (ex: "Requerente: Nome")
        elif ':' in text:
            parts = text.split(':', 1)
            label = self._strip_markdown(parts[0].strip())
            value = self._strip_markdown(parts[1].strip()) if len(parts) > 1 else ''

            # Etiqueta em negrito e maiúsculo
            run_label = p.add_run(label.upper() + ':')
            run_label.font.name = self.font_name
            run_label.font.size = Pt(self.font_size)
            run_label.bold = True

            # Valor normal
            if value:
                run_value = p.add_run(' ' + value)
                run_value.font.name = self.font_name
                run_value.font.size = Pt(self.font_size)
        else:
            # Se não tem dois pontos nem é processo, usa formatação normal
            self._add_formatted_text(p, text)
        
        # Se for o último campo (Requerido), pula uma linha após
        if is_last:
            blank_para = doc.add_paragraph()  # Linha em branco após cabeçalho
            blank_para.paragraph_format.space_before = Pt(0)
            blank_para.paragraph_format.space_after = Pt(0)
            blank_para.paragraph_format.line_spacing = 1.0
            blank_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    
    def _add_heading(self, doc: Document, text: str, level: int):
        """Adiciona título formatado - justificado, sem recuo, negrito, numerado."""
        p = doc.add_paragraph()
        
        # Remove formatação markdown do texto
        clean_text = self._strip_markdown(text)
        
        # Obtém numeração automática
        numero = self._get_heading_number(level)
        
        # Monta texto do título (com número se habilitado)
        titulo_texto = f"{numero}{clean_text.upper() if level <= 2 else clean_text}"
        
        # Configura formatação do título
        run = p.add_run(titulo_texto)
        run.font.name = self.font_name
        run.font.bold = True
        run.font.size = Pt(12)
        
        # Justificado (não centralizado), sem recuo
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.first_line_indent = Cm(0)  # SEM recuo
        p.paragraph_format.left_indent = Cm(0)  # SEM recuo esquerdo
        p.paragraph_format.line_spacing = self.line_spacing  # Espaçamento consistente
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE  # Tipo de espaçamento
        
        # Marca que o próximo parágrafo virá após um heading (para garantir formatação)
        self._last_was_heading = True
    
    def _add_paragraph(self, doc: Document, text: str, indent: bool = True):
        """Adiciona parágrafo com formatação inline e formatação consistente."""
        p = doc.add_paragraph()
        
        # Força o estilo 'Normal' explicitamente
        p.style = doc.styles['Normal']
        
        # Aplica formatação direta (sobrescreve estilo) - GARANTE formatação correta
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = self.line_spacing
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.space_after = self.space_after
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Cm(0)  # Sem recuo esquerdo
        p.paragraph_format.right_indent = Cm(0)  # Sem recuo direito
        
        # Configura recuo de primeira linha - SEMPRE aplicar explicitamente
        if indent:
            p.paragraph_format.first_line_indent = self.first_line_indent
        else:
            p.paragraph_format.first_line_indent = Cm(0)
        
        # Processa formatação inline (negrito, itálico)
        self._add_formatted_text(p, text)
        
        # Reseta flag de heading
        self._last_was_heading = False
    
    def _add_formatted_text(self, paragraph, text: str):
        """Adiciona texto com formatação inline (negrito, itálico)."""
        # Limpa marcadores órfãos antes de processar
        text = self._clean_orphan_markers(text)

        # Padrões de formatação Markdown
        # Ordem importante: negrito+itálico primeiro
        patterns = [
            (r'\*\*\*(.+?)\*\*\*', 'bold_italic'),  # ***texto***
            (r'\*\*(.+?)\*\*', 'bold'),              # **texto**
            (r'\*(.+?)\*', 'italic'),                # *texto*
            (r'_(.+?)_', 'italic'),                  # _texto_
        ]

        # Encontra todas as formatações e suas posições
        segments = []
        last_end = 0

        # Combina todos os padrões em um regex
        combined_pattern = r'(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*[^*]+?\*|_.+?_)'
        
        for match in re.finditer(combined_pattern, text):
            # Adiciona texto antes do match
            if match.start() > last_end:
                segments.append(('normal', text[last_end:match.start()]))
            
            matched_text = match.group(0)
            
            # Identifica o tipo de formatação
            if matched_text.startswith('***'):
                clean_text = matched_text[3:-3]
                segments.append(('bold_italic', clean_text))
            elif matched_text.startswith('**'):
                clean_text = matched_text[2:-2]
                segments.append(('bold', clean_text))
            elif matched_text.startswith('*') or matched_text.startswith('_'):
                clean_text = matched_text[1:-1]
                segments.append(('italic', clean_text))
            
            last_end = match.end()
        
        # Adiciona texto restante
        if last_end < len(text):
            segments.append(('normal', text[last_end:]))
        
        # Se não houver formatação, adiciona texto simples
        if not segments:
            segments = [('normal', text)]
        
        # Adiciona cada segmento ao parágrafo
        for style, segment_text in segments:
            run = paragraph.add_run(segment_text)
            run.font.name = self.font_name
            run.font.size = Pt(self.font_size)
            
            # Configura fonte para caracteres asiáticos/especiais (evita fallback)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font_name)
            
            if style == 'bold' or style == 'bold_italic':
                run.bold = True
            if style == 'italic' or style == 'bold_italic':
                run.italic = True
    
    def _add_blockquote(self, doc: Document, lines: List[str]):
        """Adiciona citação (blockquote) com formatação especial - recuo 3cm, sem itálico."""
        # Junta todas as linhas da citação
        text = ' '.join(lines)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = self.quote_indent  # 3cm
        p.paragraph_format.right_indent = Cm(0)  # Sem recuo direito
        p.paragraph_format.line_spacing = self.quote_line_spacing
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Cm(0)  # Citações sem recuo de primeira linha

        # Processa formatação inline com fonte menor (SEM itálico)
        self._add_blockquote_text(p, text)
    
    def _add_blockquote_text(self, paragraph, text: str):
        """Adiciona texto de citação com formatação (SEM itálico por padrão)."""
        # Remove asteriscos simples que envolvem todo o texto (ex: *"texto"*)
        text = re.sub(r'^\*(.+)\*$', r'\1', text.strip())
        # Remove asteriscos simples internos também (ex: *texto*)
        text = re.sub(r'\*([^*]+)\*', r'\1', text)

        # Padrão para negrito dentro da citação
        pattern = r'\*\*(.+?)\*\*'

        last_end = 0
        for match in re.finditer(pattern, text):
            # Texto antes do negrito
            if match.start() > last_end:
                run = paragraph.add_run(text[last_end:match.start()])
                run.font.name = self.font_name
                run.font.size = Pt(self.quote_font_size)
                # SEM itálico

            # Texto em negrito
            run = paragraph.add_run(match.group(1))
            run.font.name = self.font_name
            run.font.size = Pt(self.quote_font_size)
            run.bold = True
            # SEM itálico

            last_end = match.end()

        # Texto restante
        if last_end < len(text):
            run = paragraph.add_run(text[last_end:])
            run.font.name = self.font_name
            run.font.size = Pt(self.quote_font_size)
            # SEM itálico
    
    def _add_list(self, doc: Document, items: List[str], ordered: bool = False, style: str = 'number'):
        """
        Adiciona lista ordenada ou não ordenada usando numeração nativa do Word.

        Para listas alfabéticas (a, b, c...), usa numbering.xml do Word para garantir
        que a numeração seja gerenciada pelo Word e não como texto literal.

        Args:
            doc: Documento Word
            items: Lista de textos dos itens
            ordered: Se True, lista ordenada; se False, lista com bullets
            style: 'number' para 1, 2, 3... ou 'letter' para a), b), c)...
        """
        if ordered and style == 'letter':
            # Usa lista nativa do Word para letras (a, b, c...)
            self._add_letter_list_native(doc, items)
        else:
            # Para listas numéricas e bullets, usa método tradicional
            for i, item in enumerate(items):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.left_indent = self.list_indent  # Recuo de 3cm
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = self.line_spacing
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

                # Marcador
                if ordered:
                    marker = f"{i + 1}. "
                else:
                    marker = "• "

                # Adiciona marcador
                run = p.add_run(marker)
                run.font.name = self.font_name
                run.font.size = Pt(self.font_size)

                # Adiciona texto do item com formatação
                self._add_formatted_text(p, item)

    def _add_letter_list_native(self, doc: Document, items: List[str]):
        """
        Adiciona lista alfabética usando numeração NATIVA do Word (numbering.xml).

        Isso garante que a numeração a), b), c), d)... seja gerenciada pelo Word,
        evitando problemas de renderização onde todos os itens aparecem como a).

        A solução cria uma definição de numeração única para toda a lista,
        garantindo que todos os itens compartilhem o mesmo numId e a sequência
        seja mantida corretamente.
        """
        # Acessa ou cria a parte de numeração do documento
        numbering_part = doc.part.numbering_part
        if numbering_part is None:
            # Cria numbering part se não existir
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            from docx.parts.numbering import NumberingPart
            numbering_part = NumberingPart.new()
            doc.part.relate_to(numbering_part, RT.NUMBERING)

        numbering_elm = numbering_part.numbering_definitions._numbering

        # Gera IDs únicos para esta lista
        # Encontra o maior abstractNumId e numId existentes
        existing_abstract_ids = [
            int(el.get(qn('w:abstractNumId')))
            for el in numbering_elm.findall(qn('w:abstractNum'))
        ]
        existing_num_ids = [
            int(el.get(qn('w:numId')))
            for el in numbering_elm.findall(qn('w:num'))
        ]

        new_abstract_id = max(existing_abstract_ids, default=0) + 1
        new_num_id = max(existing_num_ids, default=0) + 1

        # Cria abstractNum para lista alfabética com sufixo )
        abstract_num = OxmlElement('w:abstractNum')
        abstract_num.set(qn('w:abstractNumId'), str(new_abstract_id))

        # Define o nível 0 (único nível que usamos)
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), '0')

        # Início em 1 (que corresponde a 'a')
        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)

        # Formato: lowerLetter (a, b, c...)
        num_fmt = OxmlElement('w:numFmt')
        num_fmt.set(qn('w:val'), 'lowerLetter')
        lvl.append(num_fmt)

        # Texto do nível: "%1)" para a), b), c)...
        lvl_text = OxmlElement('w:lvlText')
        lvl_text.set(qn('w:val'), '%1)')
        lvl.append(lvl_text)

        # Alinhamento à esquerda
        lvl_jc = OxmlElement('w:lvlJc')
        lvl_jc.set(qn('w:val'), 'left')
        lvl.append(lvl_jc)

        # Propriedades de parágrafo (recuo)
        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        # Recuo esquerdo de 3cm (1701 twips = 3cm) + hanging para o marcador
        ind.set(qn('w:left'), str(int(Cm(3).twips)))
        ind.set(qn('w:hanging'), '360')  # Hanging indent para o marcador
        pPr.append(ind)
        lvl.append(pPr)

        # Propriedades de fonte
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), self.font_name)
        rFonts.set(qn('w:hAnsi'), self.font_name)
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(self.font_size * 2))  # Tamanho em half-points
        rPr.append(sz)
        lvl.append(rPr)

        abstract_num.append(lvl)
        numbering_elm.append(abstract_num)

        # Cria num que referencia o abstractNum
        num = OxmlElement('w:num')
        num.set(qn('w:numId'), str(new_num_id))
        abstract_num_id = OxmlElement('w:abstractNumId')
        abstract_num_id.set(qn('w:val'), str(new_abstract_id))
        num.append(abstract_num_id)
        numbering_elm.append(num)

        # Adiciona cada item da lista usando o mesmo numId
        for item in items:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = self.line_spacing
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

            # Aplica a numeração nativa ao parágrafo
            p_pr = p._p.get_or_add_pPr()
            num_pr = OxmlElement('w:numPr')
            ilvl = OxmlElement('w:ilvl')
            ilvl.set(qn('w:val'), '0')
            num_pr.append(ilvl)
            num_id_elm = OxmlElement('w:numId')
            num_id_elm.set(qn('w:val'), str(new_num_id))
            num_pr.append(num_id_elm)
            p_pr.append(num_pr)

            # Adiciona texto do item com formatação
            self._add_formatted_text(p, item)
    
    def _strip_markdown(self, text: str) -> str:
        """Remove marcadores markdown do texto."""
        # Remove negrito
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        # Remove itálico
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'_(.+?)_', r'\1', text)
        # Remove ** órfãos (sem par)
        text = re.sub(r'\*\*', '', text)
        return text

    def _clean_orphan_markers(self, text: str) -> str:
        """
        Remove marcadores markdown órfãos (sem par) do texto.

        Exemplos de órfãos:
        - 'texto**' -> 'texto' (** no fim sem par)
        - '**texto' -> 'texto' (** no início sem par)

        Exemplos válidos (não remove):
        - '**texto**' -> mantém (par completo, será processado depois)
        """
        # Conta ocorrências de **
        count = text.count('**')

        # Se não há ** ou pares completos, retorna como está
        if count == 0 or count % 2 == 0:
            return text

        # Número ímpar de ** - há pelo menos um órfão
        # Estratégia: remove ** do final primeiro (mais comum em erros de geração)
        if text.rstrip().endswith('**'):
            text = re.sub(r'\*\*\s*$', '', text)
        elif text.lstrip().startswith('**'):
            text = re.sub(r'^\s*\*\*', '', text)
        else:
            # Órfão no meio - remove a primeira ocorrência
            text = re.sub(r'\*\*', '', text, count=1)

        return text


def create_default_template():
    """
    Cria o template DOCX padrão se não existir.
    Este template define os estilos base que serão usados nas conversões.
    """
    if TEMPLATE_PATH.exists():
        print(f"ℹ️  Template já existe: {TEMPLATE_PATH}")
        return
    
    doc = Document()
    
    # Configura margens ABNT
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
    
    # Configura estilo Normal
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.first_line_indent = Cm(2.0)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Adiciona parágrafo placeholder
    p = doc.add_paragraph("Template de Peça Jurídica - PGE-MS")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Salva template
    doc.save(str(TEMPLATE_PATH))
    print(f"✅ Template criado: {TEMPLATE_PATH}")


def markdown_to_docx(
    markdown_text: str,
    output_path: str,
    template_path: Optional[str] = None,
    **kwargs
) -> bool:
    """
    Função de conveniência para converter Markdown para DOCX.
    
    Args:
        markdown_text: Texto em formato Markdown
        output_path: Caminho de saída para o arquivo .docx
        template_path: Caminho para template personalizado (opcional)
        **kwargs: Configurações adicionais para DocxConverter
        
    Returns:
        True se conversão bem sucedida
        
    Example:
        >>> markdown_to_docx(
        ...     "# Título\\n\\nParágrafo com **negrito**.",
        ...     "output.docx",
        ...     first_line_indent_cm=2.0  # Customiza recuo
        ... )
    """
    converter = DocxConverter(template_path=template_path, **kwargs)
    return converter.convert(markdown_text, output_path)


# Cria template padrão ao importar módulo (se não existir)
if not TEMPLATE_PATH.exists():
    try:
        create_default_template()
    except Exception as e:
        print(f"⚠️ Não foi possível criar template padrão: {e}")
