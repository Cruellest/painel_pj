"""
Testes para validar a numeração de listas no conversor DOCX.

Este arquivo garante que listas alfabéticas (a, b, c, d...) sejam
geradas corretamente usando numeração nativa do Word (numbering.xml),
evitando o problema onde todos os itens aparecem como "a)".
"""

import os
import tempfile
import pytest
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

# Importa o conversor
import sys
sys.path.insert(0, str(Path(__file__).parent.parent))
from sistemas.gerador_pecas.docx_converter import DocxConverter, markdown_to_docx


class TestLetterListNumbering:
    """Testes para validar numeração de listas alfabéticas"""

    @pytest.fixture
    def converter(self):
        """Cria instância do conversor"""
        return DocxConverter(
            first_line_indent_cm=2.0,
            list_indent_cm=3.0
        )

    @pytest.fixture
    def temp_docx(self):
        """Cria arquivo temporário para saída"""
        fd, path = tempfile.mkstemp(suffix='.docx')
        os.close(fd)
        yield path
        # Cleanup
        if os.path.exists(path):
            os.remove(path)

    def test_lista_alfabetica_4_itens(self, converter, temp_docx):
        """Testa lista alfabética com 4 itens: deve gerar a), b), c), d)"""
        markdown = """## DOS PEDIDOS

a) condenação do réu ao pagamento de indenização;
b) fixação de juros moratórios desde a citação;
c) condenação em honorários advocatícios;
d) gratuidade da justiça.
"""
        result = converter.convert(markdown, temp_docx)
        assert result is True, "Conversão deve ter sucesso"

        # Abre o documento e verifica a estrutura
        doc = Document(temp_docx)

        # Encontra parágrafos com numeração
        numbered_paragraphs = []
        for p in doc.paragraphs:
            p_pr = p._p.find(qn('w:pPr'))
            if p_pr is not None:
                num_pr = p_pr.find(qn('w:numPr'))
                if num_pr is not None:
                    num_id = num_pr.find(qn('w:numId'))
                    if num_id is not None:
                        numbered_paragraphs.append({
                            'text': p.text,
                            'numId': num_id.get(qn('w:val'))
                        })

        # Deve ter 4 parágrafos numerados
        assert len(numbered_paragraphs) == 4, f"Deve ter 4 itens numerados, encontrou {len(numbered_paragraphs)}"

        # Todos devem compartilhar o mesmo numId (mesma lista)
        num_ids = set(p['numId'] for p in numbered_paragraphs)
        assert len(num_ids) == 1, f"Todos os itens devem ter o mesmo numId, encontrou: {num_ids}"

        # Verifica que o texto foi preservado (sem o marcador a), b), etc.)
        textos = [p['text'].strip() for p in numbered_paragraphs]
        assert 'condenação do réu ao pagamento de indenização' in textos[0].lower()
        assert 'juros moratórios' in textos[1].lower()
        assert 'honorários' in textos[2].lower()
        assert 'gratuidade' in textos[3].lower()

    def test_lista_alfabetica_usa_numbering_xml(self, converter, temp_docx):
        """Verifica que a lista usa numbering.xml nativo do Word"""
        markdown = """a) primeiro item;
b) segundo item;
c) terceiro item.
"""
        converter.convert(markdown, temp_docx)
        doc = Document(temp_docx)

        # Verifica que existe numbering definitions
        numbering_part = doc.part.numbering_part
        assert numbering_part is not None, "Documento deve ter numbering part"

        # Verifica que existe abstractNum com formato lowerLetter
        numbering_elm = numbering_part.numbering_definitions._numbering
        abstract_nums = numbering_elm.findall(qn('w:abstractNum'))

        found_letter_format = False
        for abstract_num in abstract_nums:
            for lvl in abstract_num.findall(qn('w:lvl')):
                num_fmt = lvl.find(qn('w:numFmt'))
                if num_fmt is not None and num_fmt.get(qn('w:val')) == 'lowerLetter':
                    found_letter_format = True
                    # Verifica também o lvlText
                    lvl_text = lvl.find(qn('w:lvlText'))
                    assert lvl_text is not None
                    assert lvl_text.get(qn('w:val')) == '%1)', "Formato deve ser %1) para a), b), c)..."
                    break

        assert found_letter_format, "Deve existir abstractNum com formato lowerLetter"

    def test_lista_numerica_nao_afetada(self, converter, temp_docx):
        """Verifica que listas numéricas (1, 2, 3) continuam funcionando"""
        markdown = """1. primeiro item
2. segundo item
3. terceiro item
"""
        result = converter.convert(markdown, temp_docx)
        assert result is True

        doc = Document(temp_docx)
        paragraphs = [p for p in doc.paragraphs if p.text.strip()]

        # Para listas numéricas, o marcador é texto
        assert '1.' in paragraphs[0].text
        assert '2.' in paragraphs[1].text
        assert '3.' in paragraphs[2].text

    def test_lista_bullets_nao_afetada(self, converter, temp_docx):
        """Verifica que listas com bullets continuam funcionando"""
        markdown = """- primeiro item
- segundo item
- terceiro item
"""
        result = converter.convert(markdown, temp_docx)
        assert result is True

        doc = Document(temp_docx)
        paragraphs = [p for p in doc.paragraphs if p.text.strip()]

        # Para bullets, o marcador é texto
        assert '•' in paragraphs[0].text
        assert '•' in paragraphs[1].text
        assert '•' in paragraphs[2].text

    def test_multiplas_listas_alfabeticas_independentes(self, converter, temp_docx):
        """Testa que múltiplas listas alfabéticas são independentes"""
        markdown = """## PRIMEIRA LISTA

a) item 1 lista 1;
b) item 2 lista 1.

## SEGUNDA LISTA

a) item 1 lista 2;
b) item 2 lista 2;
c) item 3 lista 2.
"""
        converter.convert(markdown, temp_docx)
        doc = Document(temp_docx)

        # Encontra parágrafos numerados
        numbered_paragraphs = []
        for p in doc.paragraphs:
            p_pr = p._p.find(qn('w:pPr'))
            if p_pr is not None:
                num_pr = p_pr.find(qn('w:numPr'))
                if num_pr is not None:
                    num_id = num_pr.find(qn('w:numId'))
                    if num_id is not None:
                        numbered_paragraphs.append({
                            'text': p.text,
                            'numId': num_id.get(qn('w:val'))
                        })

        # Deve ter 5 itens no total (2 + 3)
        assert len(numbered_paragraphs) == 5

        # As listas devem ter numIds diferentes
        num_ids = [p['numId'] for p in numbered_paragraphs]
        # Primeiros 2 itens: lista 1
        # Últimos 3 itens: lista 2
        assert num_ids[0] == num_ids[1], "Itens da lista 1 devem ter mesmo numId"
        assert num_ids[2] == num_ids[3] == num_ids[4], "Itens da lista 2 devem ter mesmo numId"
        assert num_ids[0] != num_ids[2], "Listas diferentes devem ter numIds diferentes"

    def test_nao_usa_texto_literal_para_marcador(self, converter, temp_docx):
        """Garante que o marcador a), b), c) NÃO é texto literal"""
        markdown = """a) primeiro;
b) segundo;
c) terceiro.
"""
        converter.convert(markdown, temp_docx)
        doc = Document(temp_docx)

        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                # O texto NÃO deve começar com "a)", "b)", "c)" como texto literal
                # O Word renderiza isso automaticamente via numbering.xml
                assert not text.startswith('a)'), f"Texto não deve começar com 'a)': {text}"
                assert not text.startswith('b)'), f"Texto não deve começar com 'b)': {text}"
                assert not text.startswith('c)'), f"Texto não deve começar com 'c)': {text}"


class TestListFormatPreservation:
    """Testes para garantir que formatação dentro dos itens é preservada"""

    @pytest.fixture
    def converter(self):
        return DocxConverter()

    @pytest.fixture
    def temp_docx(self):
        fd, path = tempfile.mkstemp(suffix='.docx')
        os.close(fd)
        yield path
        if os.path.exists(path):
            os.remove(path)

    def test_negrito_em_item_lista(self, converter, temp_docx):
        """Verifica que negrito dentro de item de lista é preservado"""
        markdown = """a) **condenação** do réu;
b) fixação de **juros**;
c) **honorários** advocatícios.
"""
        result = converter.convert(markdown, temp_docx)
        assert result is True

        doc = Document(temp_docx)

        # Verifica que existem runs com bold
        found_bold = False
        for p in doc.paragraphs:
            for run in p.runs:
                if run.bold and run.text.strip():
                    found_bold = True
                    break

        assert found_bold, "Deve preservar formatação negrito nos itens"


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
